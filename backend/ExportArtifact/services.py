import base64
import re
import markdown
from xhtml2pdf import pisa
from docx import Document
from docx.shared import Inches
import hashlib
import io

DATA_URI_IMAGE_RE = re.compile(r'^!\[[^\]]*\]\(data:image/\w+;base64,([^)]+)\)$')

def convert_to_html(md_content):
    return markdown.markdown(md_content)

def convert_to_pdf(md_content):
    html = convert_to_html(md_content)
    buffer = io.BytesIO()
    pisa.CreatePDF(html, dest=buffer)
    return buffer.getvalue()

def convert_to_docx(md_content):
    doc = Document()
    for line in md_content.split('\n'):
        image_match = DATA_URI_IMAGE_RE.match(line.strip())
        if image_match:
            doc.add_picture(io.BytesIO(base64.b64decode(image_match.group(1))), width=Inches(5.5))
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.strip():
            doc.add_paragraph(line)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def compute_hash(content):
    return hashlib.sha256(content.encode()).hexdigest()